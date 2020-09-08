import pyttsx3
import speech_recognition as sr
import pyaudio
import sys
import webbrowser as wb
from docx import Document
import smtplib
import random
import datetime
import wikipedia
from pygame import mixer
import pyowm
import re

engine = pyttsx3.init('sapi5')
voices = engine.getProperty('voices')
engine.setProperty('voice', voices[1].id)
engine.setProperty('rate', 120)


greetings = ['hey there', 'hello', 'hi', 'Hai', 'hey!', 'hey']
question = ['how are you', 'how are you doing']
responses = ['Okay', "I'm fine"]
var1 = ['who made you', 'who created you']
var2 = ['I was created by these people right in thier computer.',
        'Some guys whom i never got to know.']
var3 = ['what time is it', 'what is the time', 'time']
var4 = ['who are you', 'what is your name']
cmd1 = ['open browser', 'open Google','open google']
cmd2 = ['play music', 'play songs', 'play a song', 'open music player']
cmd3 = ['tell a joke', 'tell me a joke',
        'say something funny', 'tell something funny']
jokes = ['Can a kangaroo jump higher than a house? Of course, a house doesn’t jump at all.', 'My dog used to chase people on a bike a lot. It got so bad, finally I had to take his bike away.',
         'Doctor: Im sorry but you suffer from a terminal illness and have only 10 to live.Patient: What do you mean, 10? 10 what? Months? Weeks?!"Doctor: Nine.']
cmd4 = ['open Youtube', 'I want to watch a video']
cmd5 = ['tell me the weather', 'weather', 'what about the weather','what is the weather']
cmd6 = ['exit', 'close', 'goodbye', 'nothing']
cmd7 = ['what is your color', 'what is your colour', 'your color', 'your color?']
colrep = ['Right now its rainbow',
          'Right now its transparent', 'Right now its non chromatic']
cmd8 = ['what is you favourite colour', 'what is your favourite color']
cmd9 = ['thank you']
cmd10=['open my file','update my file','open file','open a file']

cmd11=['send email','send an email','send mail','I want to send an email']
repfr9 = ['you re welcome', 'glad i could help you']


def voice_input():
    r = sr.Recognizer()
    with sr.Microphone() as source:
        r.adjust_for_ambient_noise(source)
        print("Listening...")
        audio = r.listen(source)
        query = r.recognize_google(audio, language='en-in')
        return query


def takeCommand():
    now = datetime.datetime.now()
    r = sr.Recognizer()
    with sr.Microphone() as source:
        r.adjust_for_ambient_noise(source)
        print("Listening...")
        audio = r.listen(source)
    try:
        print("Recognizing...")
        query = r.recognize_google(audio, language='en-in')
        print(f"User Said:  {query}\n")
        
        if query in greetings:

            random_greeting = random.choice(greetings)
            print(random_greeting)
            engine.say(random_greeting)
            engine.runAndWait()
        elif query in question:

            engine.say('I am fine')
            engine.runAndWait()
            print('I am fine')
        elif query in var1:
            
            engine.runAndWait()
            reply = random.choice(var2)

            speak(reply)
        elif query in var3:

            print("Current time : ")
            print(now.strftime("The time is %H:%M"))
            engine.say(now.strftime("The time is %H:%M"))
            engine.runAndWait()
        elif query in var4:
            engine.say('I am zara,a silly bot')
            engine.runAndWait()
        elif query in cmd1:
            wb.open("https://www.google.com")
            print("opening google")
            speak("opening google")
        elif query in cmd2:
            mixer.init()
            mixer.music.load("song.wav")
            mixer.music.play()
        elif query in cmd3:
            jokrep = random.choice(jokes)
            engine.say(jokrep)
            engine.runAndWait()
        elif query in cmd4:
            wb.open("https://www.youtube.com")
            print("opening youtube")
            speak("opening youtube")
        elif query in cmd5:
            speak("which city")
            city=voice_input()
            
            
            owm = pyowm.OWM('#token')
            obs = owm.weather_at_place(city)
            w = obs.get_weather()
            k = w.get_status()
            x = w.get_temperature(unit='celsius')
            print('Current weather in %s is %s. The maximum temperature is %0.2f and the minimum temperature is %0.2f degree celcius' % (city, k, x['temp_max'], x['temp_min']))
            speak('Current weather in %s is %s. The maximum temperature is %0.2f and the minimum temperature is %0.2f degree celcius' % (city, k, x['temp_max'], x['temp_min']))
        elif query in cmd6:
            print('see you later')
            engine.say('see you later')
            engine.runAndWait()
            sys.exit(1)
        elif query in cmd7:

            print(random.choice(colrep))
            engine.say(random.choice(colrep))
            engine.runAndWait()
            print('It keeps changing every micro second')
            engine.say('It keeps changing every micro second')
            engine.runAndWait()
        elif query in cmd8:
            print(random.choice(colrep))
            engine.say(random.choice(colrep))
            engine.runAndWait()
            print('It keeps changing every micro second')
            engine.say('It keeps changing every micro second')
            engine.runAndWait()
        
        elif query in cmd9:
            print(random.choice(repfr9))
            engine.say(random.choice(repfr9))
            engine.runAndWait()
        elif query in cmd10:
            speak("which file you want to open")
            print("word file or text file")
            query=voice_input()
            print(query)
            if query=="text file":

                speak("what is the name of file")
                query=voice_input()
                query=query+'.txt'
                file=open(query,'w')
                speak("what do you want to write")
                content=voice_input()
                file.write(content) 
                file.close() 
                wb.open("C:/Users/Richa/Desktop/New folder/"+query)
                print("your file is updated.")
            else:
                
                speak("what is the name of file")
                command=voice_input()
                command=command+'.docx'
                document = Document()
                speak("what is the heading")
                heading=voice_input()
                document.add_heading(heading, 0)
                speak("what do you want to write")
                para=voice_input()
                document.add_paragraph(para)
                document.add_page_break()

                document.save(command)
                wb.open("C:/Users/Richa/Desktop/New folder/"+command)
                print("your file is updated.")
        elif query in cmd11:
            # creates SMTP session 
            s = smtplib.SMTP('smtp.gmail.com', 587) 
    
            # start TLS for security 
            s.starttls() 
  
            # Authentication 
            s.login("senderid", "senderpsw") 
            #recipient name
            speak("whom do you want to send email")
            r=voice_input()
            r=r.replace(" ","")
            
            recipient=r+"@gmail.com"

            # message to be sent 
            speak("what message you want to send")
            message = voice_input()
            #confirmation
            speak("is it correct")
            print(recipient)
            print(message)
            t=voice_input()
            if t=="yes":

                # sending the mail 
                s.sendmail("richa75937@gmail.com", recipient, message) 
                 # terminating the session 
                s.quit() 

                print("email successfully sent")
                speak("email successfully sent")    
            
            else:
                print("sorry!!")    
                speak("try again")
                # terminating the session 
                s.quit() 
        

        else:
            engine.say("please wait")
            engine.runAndWait()
            print(wikipedia.summary(r.recognize_google(audio)))
            engine.say(wikipedia.summary(r.recognize_google(audio)))
            engine.runAndWait()
            userInput3 = input("or else search in google")
            wb.open_new('www.google.com/search?q=' + userInput3)

    except Exception :
        print("Say That Again please...")
        speak("Say That Again please...")
        return "None"


def speak(audio):
    engine.say(audio)
    engine.runAndWait()


if __name__ == "__main__":
    print("Hey!!I am Zara.")
    speak("hey i m zara")
    while True:
        takeCommand()
